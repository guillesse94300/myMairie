# -*- coding: utf-8 -*-
"""Génère un résumé Word des mentions trouvées dans les PDF, avec paragraphes et détails."""
import re
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOSSIER = Path(__file__).resolve().parent
TERMES = [
    "Bois d'Haucourt",
    "Sente Brunehaut",
    "Vertefeuille",
    "VTT",
    "Chasse",
]

# Contexte max autour du terme (caractères avant/après) si pas de paragraphe clair
TAILLE_CONTEXTE = 400

def normalise(s):
    return s.replace("\r", " ").replace("\n", " ").replace("  ", " ")

def nettoyer_paragraphe(s):
    """Retire espaces superflus, garde le texte lisible."""
    if not s or not s.strip():
        return ""
    return " ".join(s.split()).strip()

def extraire_extraits(texte_page, terme, num_page):
    """Retourne la liste des extraits (paragraphes ou contexte) contenant le terme."""
    extraits = []
    terme_lower = terme.lower()
    texte_lower = texte_page.lower()

    if terme_lower not in texte_lower:
        return extraits

    # 1) Découper en paragraphes (séparateurs : double newline ou double espace)
    blocs = re.split(r"\n\s*\n", texte_page)

    for bloc in blocs:
        bloc = nettoyer_paragraphe(bloc)
        if len(bloc) < 3:
            continue
        if terme_lower in bloc.lower():
            if bloc not in [e[1] for e in extraits]:
                extraits.append((num_page, bloc))

    # 2) Si aucun paragraphe trouvé, utiliser un contexte autour du terme
    if not extraits:
        idx = texte_page.lower().find(terme_lower)
        if idx != -1:
            debut = max(0, idx - TAILLE_CONTEXTE)
            fin = min(len(texte_page), idx + len(terme) + TAILLE_CONTEXTE)
            contexte = nettoyer_paragraphe(texte_page[debut:fin])
            if contexte:
                extraits.append((num_page, contexte))

    return extraits

def main():
    # resultats[terme] = [(nom_fichier, [(num_page, extrait), ...]), ...]
    resultats = {t: [] for t in TERMES}
    pdfs = sorted(p for p in DOSSIER.glob("*.pdf") if p.is_file())

    for path in pdfs:
        try:
            reader = PdfReader(str(path))
            for terme in TERMES:
                details_fichier = []
                for num_page, page in enumerate(reader.pages, start=1):
                    texte_page = page.extract_text() or ""
                    for num_p, extrait in extraire_extraits(texte_page, terme, num_page):
                        details_fichier.append((num_p, extrait))
                if details_fichier:
                    # Dédupliquer les extraits identiques
                    vus = set()
                    unique = []
                    for num_p, ext in details_fichier:
                        cle = (num_p, ext[:200])
                        if cle not in vus:
                            vus.add(cle)
                            unique.append((num_p, ext))
                    resultats[terme].append((path.name, unique))
        except Exception as e:
            print(f"Erreur {path.name}: {e}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titre = doc.add_heading("Résumé des mentions dans les documents", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "Recherche effectuée dans l’ensemble des PDF du dossier Mairie "
        "(procès-verbaux et comptes-rendus de conseil municipal). "
        "Pour chaque terme, sont indiqués les fichiers, les numéros de page et les paragraphes contenant la mention."
    )
    doc.add_paragraph()

    for terme in TERMES:
        doc.add_heading(terme, level=1)
        entrees = resultats[terme]
        if entrees:
            for nom_fichier, extraits in entrees:
                p_fichier = doc.add_paragraph()
                p_fichier.add_run(f"Fichier : ").bold = True
                p_fichier.add_run(nom_fichier)
                for num_page, extrait in extraits:
                    p_page = doc.add_paragraph()
                    p_page.add_run(f"Page {num_page} : ").bold = True
                    p_page.add_run(extrait)
                doc.add_paragraph()
            p_total = doc.add_paragraph()
            p_total.add_run(f"Total : {len(entrees)} fichier(s).").italic = True
        else:
            doc.add_paragraph("Aucune mention trouvée.")
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph("— Document généré automatiquement à partir de la recherche dans les PDF. —").alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = DOSSIER / "Resume-mentions-documents.docx"
    doc.save(str(out_path))
    print(f"Document enregistré : {out_path}")

if __name__ == "__main__":
    main()
