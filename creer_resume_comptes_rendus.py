# -*- coding: utf-8 -*-
"""Génère un résumé de 2 pages des comptes rendus du conseil municipal de Pierrefonds."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOSSIER = Path(__file__).resolve().parent
OUT_PATH = DOSSIER / "Resume_comptes_rendus_conseil_municipal.docx"


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Titre
    titre = doc.add_heading("Résumé des comptes rendus du conseil municipal", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Commune de Pierrefonds — Synthèse des délibérations et débats (juin 2024 – novembre 2025)")
    doc.add_paragraph()

    # === PAGE 1 ===
    doc.add_heading("1. Période et cadre des séances", level=1)
    doc.add_paragraph(
        "Les comptes rendus couvrent les séances du conseil municipal de juin 2024 à novembre 2025. "
        "Les séances se tiennent en mairie sous la présidence de Mme Florence DEMOUY, maire. "
        "Les délibérations portent sur des décisions prises en séance ou par délégation au maire (DM). "
        "Les thèmes récurrents sont les finances, l’enfance et la vie scolaire, les travaux et la voirie, "
        "l’intercommunalité, le personnel et les associations."
    )

    doc.add_heading("2. Finances et budget", level=1)
    doc.add_paragraph(
        "Plusieurs décisions budgétaires modificatives (DBM) sont adoptées pour ajuster les crédits "
        "(rue de l’Armistice, provisions, régularisations). Le compte financier unique (CFU) 2024 est approuvé ; "
        "l’affectation des résultats et le budget primitif 2025 sont votés (section fonctionnement environ 1,58 M€, "
        "investissement environ 1,48 M€). Les taux d’imposition 2025 sont fixés (taxe foncière bâties 40,92 %, "
        "non bâties 43,95 %, taxe d’habitation 9,97 % pour les résidences secondaires). Des provisions sont constituées "
        "(ex. créance La Pétrifontaine 648 €). Des admissions en non-valeur sont prononcées pour créances irrécouvrables. "
        "Les redevances d’occupation du domaine public (RODP télécom, électricité, chantiers provisoires) sont fixées chaque année. "
        "Une tarification annuelle (emprises commerciales, marché, Foyer Napoléon, cimetière, bibliothèque, etc.) est votée "
        "et reconduite (gratuité bibliothèque pour les moins de 18 ans). La participation à la prévoyance et à la "
        "protection sociale complémentaire des agents (labellisation) est mise en place ou ajustée."
    )

    doc.add_heading("3. Enfance, scolaire et périscolaire", level=1)
    doc.add_paragraph(
        "L’organisation du temps scolaire en 4 jours est reconduite (dérogation triennale). Les tarifs de la restauration "
        "scolaire, du périscolaire (matin/soir) et des accueils de loisirs sont fixés selon le quotient familial CAF et "
        "révisés chaque année (repas cantine, goûter, majoration en cas d’inscription tardive). La participation des "
        "communes de résidence pour les enfants accueillis en classe ULIS est fixée (350 € puis 700 € par élève). "
        "Le marché de fourniture de repas est attribué à Dupont Restauration (repas unitaire 2,82 € HT). Des contrats "
        "d’engagement éducatif (CEE) sont créés pour les accueils de loisirs (animateurs, directeurs). La commune participe "
        "au financement d’une classe découverte (CP, thème cirque). Des questions récurrentes portent sur les impayés "
        "cantine et le suivi des familles en difficulté."
    )

    doc.add_heading("4. Travaux, voirie et aménagements", level=1)
    doc.add_paragraph(
        "Les travaux de la rue de l’Armistice (RD 973) sont menés en trois phases : phase 1 (camping – Clos Saint-Ladre) "
        "achevée, phase 2 (trottoirs jusqu’à la sente de la Sautelle) partiellement réalisée avec décalage, phase 3 "
        "(sente de la Sautelle – carrefour Zwingenberg) programmée (mini-giratoire, trottoirs, potelets PMR). Le coût "
        "phase 3 est estimé à environ 344 k€ HT, avec subventions État (DETR), Conseil départemental et Région (FAPL). "
        "Une place de dépôt ONF en forêt avec emprise sur chemin communal est autorisée. La réhabilitation de la fontaine "
        "place de l’hôtel de ville est engagée (devis Hydrogénie, végétalisation, bacs, candélabres) ; l’ABF est saisi. "
        "Les aménagements de la place (bacs composite, bancs) font débat (esthétique, cohérence patrimoniale). "
        "Le giratoire Zwingenberg est discuté (plantation d’arbre ou sculpture, essences, sécurité). La voie verte "
        "Pierrefonds–Palesne est évoquée (liaison mobilité douce vers Compiègne et Villers-Cotterêts). Une amende "
        "administrative pour dépôts sauvages est instaurée (barème selon volume). Zone bleue et stationnement sont évoqués."
    )

    # === PAGE 2 ===
    doc.add_heading("5. Intercommunalité", level=1)
    doc.add_paragraph(
        "La commune adhère au groupement d’achat d’énergies coordonné par le SE60 (Syndicat d’énergie de l’Oise). "
        "Elle approuve l’adhésion de Montmartin au SIVOC Atelier musical de l’Oise. Les rapports sur le prix et la "
        "qualité des services (RPQS) eau potable, assainissement collectif et non collectif, déchets – CCLO – sont "
        "présentés. La dissolution du SMIOCE et les écritures comptables associées donnent lieu à une DBM. La commune "
        "est actionnaire de la SPL ADTO-SAO ; le rapport d’observations définitives de la Chambre régionale des comptes "
        "(exercices 2018–2023) est présenté et pris en compte (rappel au droit sur les statuts, recommandations sur le "
        "contrôle et le pilotage). Le rapport annuel ADTO-SAO est approuvé."
    )

    doc.add_heading("6. Patrimoine, tourisme et vie locale", level=1)
    doc.add_paragraph(
        "Une convention avec le CMN (Centre des monuments nationaux) prévoit un petit train touristique desservant le "
        "château (30 % des recettes reversés à la commune) ; le parcours est modifié (départ derrière la mairie, passage "
        "rue Sabatier, sœur Aurélie, bois d’Haucourt, gare, château). Le kiosque MEDIAKIOSK fait l’objet d’une convention "
        "(affichage, pas de commerce de presse, casiers vélos/motos envisagés). La cérémonie des vœux a lieu au château. "
        "Le dossier du Commerce (incendie 07/08/2023, SNC Le tabac du château, SCI Les Coquelicots) est suivi : conflit "
        "propriétaire–exploitants, permis de construire, ERP ; la mairie indique accompagner sans pouvoir trancher un "
        "litige privé. Un festival de musique est organisé par l’association L’Enceinte au stade et parc du château ; "
        "des retours sur les nuisances sonores et le lieu sont discutés, ainsi que le prêt de véhicules municipaux aux associations."
    )

    doc.add_heading("7. Affaires générales et cadre de vie", level=1)
    doc.add_paragraph(
        "L’Association foncière de Pierrefonds est dissoute ; les équipements issus du remembrement sont incorporés au "
        "patrimoine communal. Une convention avec l’association Cats in The Air 2016 est signée pour la gestion des "
        "chats errants (capture, stérilisation, identification, relâcher) avec participation aux frais vétérinaires "
        "(plafond 1 000 €). Adhésion au CAUE 60 pour l’accompagnement des projets (cour d’école, budget participatif). "
        "Convention de partenariat avec l’Institut Charles Quentin (Saint-Louis Poissy) pour l’entretien à visée "
        "pédagogique d’espaces publics ; subvention 400 €. Cession de bail rural (famille Beguin), concessions funéraires, "
        "convention SIVOC ancienne PMI, location place de dépôt ONF pour travaux rue Armistice. Désherbeur cédé à "
        "l’association Un château pour l’Emploi. Renouvellement contrat Illiwap, logiciels métiers (JVS MAIRISTEM)."
    )

    doc.add_heading("8. Personnel et organisation", level=1)
    doc.add_paragraph(
        "Autorisations spéciales d’absence pour événements familiaux et de la vie courante (mariage, décès, garde d’enfant "
        "malade, déménagement, concours, rentrée scolaire, etc.) sont fixées après avis du comité social territorial. "
        "Création de postes vacataires (animation, technique) et d’emplois permanents (adjoints techniques 30 h et 22 h). "
        "Participation au financement de la protection sociale complémentaire (prévoyance, santé) dans le cadre de la "
        "labellisation. Adhésion au contrat d’assurance des risques statutaires du CDG 60. Cartes cadeaux pour les "
        "enfants du personnel à Noël (50 €)."
    )

    doc.add_heading("9. Subventions aux associations", level=1)
    doc.add_paragraph(
        "Les subventions annuelles sont votées (UNC, Bols d’air, Compagnie d’arc, Coopérative scolaire, Amis des Petites "
        "Bouilles, Festival des forêts, L’Enceinte, Amitié Loisirs Ainés, Judo, APE, Comité des Fêtes, Vital’Marche, "
        "P.E.P.I.T.E.S., Sauvegarde du patrimoine des forêts du Compiégnois, etc.). Les montants et critères (animation "
        "locale, ouverture au public) donnent lieu à des débats et parfois des abstentions ou votes contre (Festival des "
        "forêts, L’Enceinte, P.E.P.I.T.E.S.)."
    )

    doc.add_heading("10. Débats et questions diverses", level=1)
    doc.add_paragraph(
        "Plusieurs sujets reviennent en séance : coût réel de la zone bleue (main-d’œuvre, matériel), publication des "
        "comptes rendus (Illiwap), tribunes dans le journal municipal, demandes de « questions diverses » en fin de "
        "conseil. Des tensions apparaissent entre majorité et opposition (départs de séance, votes contre ou abstentions). "
        "Des précisions sont demandées sur la vidéoprotection, les atteintes aux biens, le plan communal de sauvegarde, "
        "la rénovation énergétique des bâtiments, l’eau et les investissements CCLO, les conventions de servitude, la "
        "santé et le remplacement du médecin. Le procès-verbal est enregistré ; les convocations et ordres du jour "
        "respectent le règlement intérieur."
    )

    doc.add_paragraph()
    p_fin = doc.add_paragraph(
        "— Document de synthèse établi à partir des comptes rendus du conseil municipal de Pierrefonds (base vectorielle et PDF). —"
    )
    p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT_PATH))
    print(f"Résumé enregistré : {OUT_PATH}")


if __name__ == "__main__":
    main()
